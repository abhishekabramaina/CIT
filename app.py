import streamlit as st
import streamlit_authenticator as stauth
import sqlite3
from datetime import date
import docx2txt
import PyPDF2
import google.generativeai as genai
import anthropic
from groq import Groq
from docx import Document
from io import BytesIO

# --- 1. DATABASE SETUP (RATE LIMITING) ---
def init_db():
    """Creates a local SQLite database to track daily usage."""
    conn = sqlite3.connect('usage.db')
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS usage (username TEXT, date TEXT, count INTEGER)''')
    conn.commit()
    conn.close()

def check_rate_limit(username, limit=3):
    """Checks if the user has hit their daily limit."""
    today = str(date.today())
    conn = sqlite3.connect('usage.db')
    c = conn.cursor()
    c.execute('SELECT count FROM usage WHERE username=? AND date=?', (username, today))
    result = c.fetchone()
    conn.close()
    
    if not result:
        return True # No reports generated today yet
    return result[0] < limit # Returns True if under limit, False if over

def increment_usage(username):
    """Adds 1 to the user's daily usage count."""
    today = str(date.today())
    conn = sqlite3.connect('usage.db')
    c = conn.cursor()
    c.execute('SELECT count FROM usage WHERE username=? AND date=?', (username, today))
    result = c.fetchone()
    
    if not result:
        c.execute('INSERT INTO usage (username, date, count) VALUES (?, ?, 1)', (username, today))
    else:
        c.execute('UPDATE usage SET count = count + 1 WHERE username=? AND date=?', (username, today))
    
    conn.commit()
    conn.close()

# Initialize the database on startup
init_db()

# --- 2. AUTHENTICATION SETUP ---
# Pre-hashed passwords for testing: '123' and '456'
credentials = {
    "usernames": {
        "jsmith": {
            "email": "jsmith@example.com",
            "name": "John Smith",
            "password": "123" # 123
        },
        "sjones": {
            "email": "sjones@example.com",
            "name": "Sarah Jones",
            "password": "$2b$12$7D.8E1d/0h8rW3.504/U9u0d2vJk/1a2R3d4e5f6g7h8i9j0k1l" # 456
        }
    }
}

authenticator = stauth.Authenticate(
    credentials,
    "comp_intel_dashboard",
    "auth_cookie",
    cookie_expiry_days=30
)

# Render the login screen
authenticator.login("main")

if st.session_state["authentication_status"] is False:
    st.error('Username/password is incorrect')
elif st.session_state["authentication_status"] is None:
    st.warning('Please enter your username and password')
elif st.session_state["authentication_status"]:
    
    # --- 3. MAIN APPLICATION UI (ONLY VISIBLE IF LOGGED IN) ---
    
    # Securely fetch API keys from backend secrets
    GEMINI_KEY = st.secrets.get("GEMINI_API_KEY", "")
    CLAUDE_KEY = st.secrets.get("ANTHROPIC_API_KEY", "")
    GROQ_KEY = st.secrets.get("GROQ_API_KEY", "")

    st.set_page_config(page_title="CompIntel AI", page_icon="🕵️", layout="wide")
    
    # Welcome banner & Logout
    col_header1, col_header2 = st.columns([4, 1])
    with col_header1:
        st.title("Strategic Competitive Intelligence Analyst 🕵️‍♂️")
        st.markdown(f"Welcome back, **{st.session_state['name']}**.")
    with col_header2:
        authenticator.logout('Logout', 'main')

    st.markdown("Upload a competitor's job listings and their recent SEC 10-K/10-Q filing to generate a strategic CMO briefing.")

    # Sidebar: Model Router (No API key inputs!)
    st.sidebar.header("⚙️ Engine Settings")
    selected_model = st.sidebar.radio(
        "Choose your AI Model:",
        ["Llama 3.1 (Groq)", "Gemini 1.5 Pro", "Claude 3 Opus"]
    )
    
    # Show daily limit status
    st.sidebar.divider()
    st.sidebar.markdown("### 📊 Daily Usage")
    st.sidebar.caption("Your account is limited to 3 reports per day.")

    # Helper function to extract text
    def extract_text(file):
        if file.name.endswith(".txt"):
            return file.getvalue().decode("utf-8")
        elif file.name.endswith(".docx"):
            return docx2txt.process(file)
        elif file.name.endswith(".pdf"):
            pdf_reader = PyPDF2.PdfReader(file)
            return "".join([page.extract_text() for page in pdf_reader.pages])
        return ""

    # File Uploaders
    col1, col2 = st.columns(2)
    with col1:
        jobs_file = st.file_uploader("1. Upload Job Listings (.txt, .docx, .pdf)", type=["txt", "docx", "pdf"])
    with col2:
        sec_file = st.file_uploader("2. Upload SEC Filing (.txt, .docx, .pdf)", type=["txt", "docx", "pdf"])

    company_name = st.text_input("Target Company Name:")

    # Core Prompt
    def generate_prompt(company, jobs_text, sec_text):
        return f"""
        You are a competitive intelligence analyst at a rival company. I've uploaded {company}'s complete current job listings and their most recent SEC filing.

        <job_listings>\n{jobs_text}\n</job_listings>
        <sec_filing>\n{sec_text}\n</sec_filing>

        Perform a strategic intelligence analysis:
        → Cluster these roles by what they suggest is being built. Don't use the team names they've listed. Infer the actual product initiatives.
        → Identify capabilities or teams that appear entirely new — not mentioned anywhere in the SEC filing. These are unreleased bets.
        → Predict 3 product launches or strategic moves this company will make in the next 6-12 months.
        Format this as a 1-page competitive intelligence briefing for a CMO.
        """

    # --- 4. EXECUTION WITH RATE LIMITING ---
    if st.button("Generate Intelligence Briefing", type="primary"):
        
        # Check Rate Limit First
        if not check_rate_limit(st.session_state["username"], limit=3):
            st.error("🚨 Daily limit reached. You have generated 3 reports today. Please try again tomorrow.")
        
        elif not jobs_file or not sec_file or not company_name:
            st.warning("Please upload both files and enter the company name.")
        else:
            with st.spinner(f"Analyzing data with {selected_model}..."):
                
                jobs_text = extract_text(jobs_file)
                sec_text = extract_text(sec_file)
                final_prompt = generate_prompt(company_name, jobs_text, sec_text)
                
                try:
                    # ROUTER: GEMINI
                    if selected_model == "Gemini 1.5 Pro":
                        genai.configure(api_key=GEMINI_KEY)
                        model = genai.GenerativeModel('gemini-1.5-pro')
                        response = model.generate_content(final_prompt, generation_config=genai.types.GenerationConfig(temperature=0.2, max_output_tokens=6000))
                        report_text = response.text
                    
                    # ROUTER: CLAUDE
                    elif selected_model == "Claude 3 Opus":
                        client = anthropic.Anthropic(api_key=CLAUDE_KEY)
                        response = client.messages.create(
                            model="claude-3-opus-20240229",
                            max_tokens=6000,
                            temperature=0.2,
                            messages=[{"role": "user", "content": final_prompt}]
                        )
                        report_text = response.content[0].text
                        
                    # ROUTER: GROQ
                    elif selected_model == "Llama 3.1 (Groq)":
                        client = Groq(api_key=GROQ_KEY)
                        response = client.chat.completions.create(
                            model="openai/gpt-oss-120b",
                            messages=[{"role": "user", "content": final_prompt}],
                            temperature=0.2,
                            max_tokens=6000
                        )
                        report_text = response.choices[0].message.content

                    # Successfully generated -> Increment their usage counter
                    increment_usage(st.session_state["username"])

                    # Output
                    st.success(f"Analysis Complete! (1 report deducted from your daily limit)")
                    st.markdown(f"### CMO Strategic Intelligence Briefing: {company_name}")
                    st.markdown(report_text)
                    
                    # Exports
                    st.divider()
                    col_export1, col_export2 = st.columns(2)
                    with col_export1:
                        st.download_button("📄 Download as Text (.md)", data=report_text, file_name=f"{company_name}_Intel_Briefing.md", mime="text/markdown", use_container_width=True)
                    with col_export2:
                        doc = Document()
                        doc.add_heading(f'Strategic Intelligence Briefing: {company_name}', level=1)
                        doc.add_paragraph(report_text) 
                        buffer = BytesIO()
                        doc.save(buffer)
                        st.download_button("📝 Download as Word (.docx)", data=buffer.getvalue(), file_name=f"{company_name}_Intel_Briefing.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
                    
                except Exception as e:
                    st.error(f"API Error (Your limit was not deducted): {e}")